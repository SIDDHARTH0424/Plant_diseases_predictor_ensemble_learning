from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def setup_document():
    doc = Document()
    
    # Title
    title = doc.add_heading("Plant Disease Detection System", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The Emseble Plant Clinic is an intelligent, research-grade plant disease detection "
        "application. It leverages a rigorous ensemble of advanced Vision Transformers (ViT) "
        "and Convolutional Neural Networks (EfficientNetB4, ResNet50) to achieve high-accuracy "
        "disease identification across numerous plant species. It features a modern, interactive web "
        "application that takes advantage of computer vision, OOD (Out of Distribution) checks, and "
        "MC Dropout for uncertainty estimation."
    )
    
    doc.add_heading("2. Architecture & Pipeline", level=1)
    doc.add_heading("2.1 Model Architecture", level=2)
    doc.add_paragraph(
        "The core AI engine is a weighted ensemble model consisting of three state-of-the-art "
        "architectures:\n"
        "- EfficientNetB4\n"
        "- Vision Transformer (ViT-B/16)\n"
        "- ResNet50\n"
        "We compute the weighted average predictions of these models to maximize robustness."
    )
    doc.add_heading("2.2 Inference Pipeline", level=2)
    doc.add_paragraph(
        "The inference pipeline employs the following steps:\n"
        "1. Out-of-Distribution (OOD) Check: Validates that the input image strictly contains plant material utilizing color-channel dominance thresholds.\n"
        "2. YOLO Leaf Detection: Isolates the primary leaf bounding box for noise-reduction.\n"
        "3. Ensemble Classification: Processes the image through Test-Time Augmentation (TTA) and Monte Carlo (MC) Dropout (10 passes) for highly calibrated classification and uncertainty estimation.\n"
        "4. Knowledge Base Retrieval: Looks up severity, immediate action, organic options, and product recommendations."
    )
    
    doc.add_heading("3. Dataset & Preprocessing", level=1)
    doc.add_paragraph(
        "The unified training dataset was constructed by extensively merging, deduplicating, and normalizing three major open-source datasets:\n"
        "- PlantDoc\n"
        "- FieldPlant\n"
        "- PlantSegV2\n"
        "All data is standardized with rigid splits across diverse pathogen profiles."
    )

    doc.add_heading("4. Evaluation & Results", level=1)
    doc.add_paragraph(
        "The ensemble model was rigorously tested over a held-out test set comprising 2,623 images. "
        "TTA (Test Time Augmentation) and MC Dropout were utilized for optimal performance."
    )
    
    # Try to add table, standard styles
    try:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1' # Try a standard Word table style
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Precision'
        hdr_cells[2].text = 'Recall'
        hdr_cells[3].text = 'F1-Score'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Macro Average'
        row_cells[1].text = '80.69%'
        row_cells[2].text = '79.96%'
        row_cells[3].text = '79.65%'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Weighted Average'
        row_cells[1].text = '84.38%'
        row_cells[2].text = '83.64%'
        row_cells[3].text = '83.60%'
    except Exception:
        # Fallback if style missing
        pass

    doc.add_paragraph("\nOverall Accuracy: 83.64% (Highest accuracy observed among healthy baselines and robust classifiers like Cassava Mosaic yielding 96.98% F1-score).")

    doc.add_heading("5. Web Interface (Frontend & API)", level=1)
    doc.add_paragraph(
        "A beautiful, responsive web interface was constructed utilizing HTML, sleek CSS glassmorphism styling, and Vanilla JavaScript. "
        "It supports both file uploads and live camera captures via navigator.mediaDevices. The backend is robustly powered by FastAPI, dynamically returning curated diagnostics over a unified port."
    )

    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "The system achieves highly scalable, accurate, and easily accessible plant diagnostics, making it an indispensable tool for farmers, researchers, and hobbyists. The API endpoints and responsive interface facilitate seamless future integration into mobile platforms or offline environments."
    )

    doc.save("Emseble_Project_Report.docx")
    print("Report generated successfully as Emseble_Project_Report.docx")

if __name__ == "__main__":
    setup_document()
